#!/usr/bin/env python3
"""
LogGuard Stress Test — Word Report Generator
=============================================
Reads the JSON output from stress_test.py and produces a formatted .docx report.

Usage:
    pip install python-docx
    python3 generate_stress_report.py stress_test_results_<timestamp>.json

    # Or auto-pick the most recent results file:
    python3 generate_stress_report.py

Output:
    stress_test_report_<timestamp>.docx  (same directory as the JSON file)
"""

import os
import sys
import json
import glob
import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Run:  pip install python-docx")
    sys.exit(1)

# ─────────────────────────────────────────────────────────────────────────────
# Colour palette
# Hex strings (RRGGBB) for cell backgrounds; RGBColor for font colours only
# ─────────────────────────────────────────────────────────────────────────────
BG_HEADER   = "1F497D"   # dark navy  — cell background
BG_ALT_ROW  = "D9E2F3"   # light blue — cell background

COLOR_HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)   # white      — font
COLOR_CRASH      = RGBColor(0xC0, 0x00, 0x00)   # dark red   — font
COLOR_OK         = RGBColor(0x37, 0x86, 0x10)   # dark green — font
COLOR_ACCENT     = RGBColor(0x1F, 0x49, 0x7D)   # navy       — font (headings)
COLOR_SUBHEADING = RGBColor(0x2E, 0x74, 0xB5)   # mid blue   — font

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour. hex_color = 'RRGGBB' string."""
    tc   = cell._tc
    tc_properties = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_properties.append(shd)

def _set_cell_border(cell, border_sides=("top","bottom","left","right"), size=4, color="1F497D"):
    """Add border to specific sides of a cell."""
    tc   = cell._tc
    tc_properties = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in border_sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_properties.append(borders)

def _bold_run(para, text, size=None, color=None):
    run = para.add_run(text)
    run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def _add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = COLOR_ACCENT if level == 1 else COLOR_SUBHEADING
    run.font.size = Pt(14) if level == 1 else Pt(12)
    return para

def _add_kv(doc, key, value, value_color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    k = para.add_run(f"{key}: ")
    k.bold = True
    k.font.size = Pt(10)
    v = para.add_run(str(value))
    v.font.size = Pt(10)
    if value_color:
        v.font.color.rgb = value_color
    return para

def _add_table(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        _set_cell_bg(cell, BG_HEADER)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run(h)
        run.bold = True
        run.font.color.rgb = COLOR_HEADER_FG
        run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = BG_ALT_ROW if r_idx % 2 == 0 else None
        for c_idx, val in enumerate(row_data):
            cell = row_cells[c_idx]
            if bg:
                _set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run  = para.add_run(str(val))
            run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table

# ─────────────────────────────────────────────────────────────────────────────
# Build the Word document
# ─────────────────────────────────────────────────────────────────────────────
def build_report(data: dict, out_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("LogGuard Server Stress Test Report")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = COLOR_ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    sr = sub.add_run(f"Generated: {datetime.datetime.now(datetime.timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)  # grey — font colour, RGBColor is fine here

    doc.add_paragraph()

    # ── 1. Test Overview ───────────────────────────────────────────────────
    _add_heading(doc, "1. Test Overview")

    _add_kv(doc, "Test Run At",        data.get("test_run_at", "—"))
    _add_kv(doc, "LogGuard Endpoint",  data.get("logguard_endpoint", "—"))
    _add_kv(doc, "Organisations",      ", ".join(data.get("organisations_tested", [])))
    _add_kv(doc, "Batch Size",         f"{data.get('batch_size', '—')} log lines per POST request")
    _add_kv(doc, "Total Elapsed",      f"{data.get('total_elapsed_seconds', '—')} seconds")

    doc.add_paragraph()

    # ── 2. Results Summary ─────────────────────────────────────────────────
    _add_heading(doc, "2. Results Summary")

    crash = data.get("crash_detected", False)
    crash_color = COLOR_CRASH if crash else COLOR_OK
    crash_text  = "CRASH DETECTED" if crash else "No crash — server survived all phases"

    _add_kv(doc, "Outcome",              crash_text, value_color=crash_color)
    _add_kv(doc, "Total Logs Accepted",  f"{data.get('total_logs_accepted', 0):,}")
    _add_kv(doc, "Total POST Requests",  f"{data.get('total_post_requests', 0):,}")
    _add_kv(doc, "POST Error Rate",      f"{data.get('post_error_rate_pct', 0):.2f}%")
    _add_kv(doc, "Peak Rate (per org)",  f"{data.get('peak_rate_per_org', '—')} logs/s")
    _add_kv(doc, "Peak Rate (total)",    f"{data.get('peak_rate_total_logs_per_s', '—')} logs/s across both orgs")
    _add_kv(doc, "LogGuard Baseline",    f"{data.get('logguard_baseline_ms', '—')} ms (measured at start)")
    _add_kv(doc, "LogGuard Final Avg",   f"{data.get('logguard_avg_ms_at_end', '—')} ms (last 5 probes)")

    doc.add_paragraph()

    # ── 3. Crash Analysis ──────────────────────────────────────────────────
    _add_heading(doc, "3. Crash Analysis")

    if crash:
        _add_kv(doc, "Crash Detected",     "YES", value_color=COLOR_CRASH)
        _add_kv(doc, "Crash Time",         data.get("crash_time", "—"))
        _add_kv(doc, "Crash Reason",       data.get("crash_reason", "—"), value_color=COLOR_CRASH)
        _add_kv(doc, "Logs at Crash",      f"{data.get('crash_at_total_logs', 0):,} total logs accepted when crash declared")

        doc.add_paragraph()
        note = doc.add_paragraph()
        note.paragraph_format.left_indent = Inches(0.3)
        nr = note.add_run(
            "Recovery steps: SSH into the K8s node and run:\n"
            "  kubectl get pods -n logguard\n"
            "  kubectl rollout restart deployment/backend -n logguard\n"
            "  kubectl rollout restart deployment/anomaly-detection -n logguard\n"
            "  kubectl rollout restart deployment/elasticsearch -n logguard\n"
            "Kubernetes will auto-restart crashed pods within ~30 seconds."
        )
        nr.font.size = Pt(9)
        nr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    else:
        _add_kv(doc, "Crash Detected", "No — server handled all phases successfully", value_color=COLOR_OK)
        note = doc.add_paragraph(
            "The server did not crash during this test run. "
            "Consider increasing the peak rate or duration to find the breaking point."
        )
        note.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── 4. Phase-by-Phase Breakdown ────────────────────────────────────────
    _add_heading(doc, "4. Phase-by-Phase Breakdown")

    phases = data.get("phases_run", [])
    if phases:
        headers = ["Phase", "Rate/org\n(logs/s)", "Total rate\n(logs/s)", "Duration\n(s)", "Logs sent", "Throughput\n(logs/s)"]
        rows    = [
            [
                p.get("phase", "—"),
                p.get("rate_per_org", "—"),
                p.get("total_rate", "—"),
                p.get("duration_s", "—"),
                f"{p.get('logs_sent', 0):,}",
                p.get("throughput_logs_s", "—"),
            ]
            for p in phases
        ]
        _add_table(doc, headers, rows, col_widths=[1.1, 1.0, 1.0, 0.9, 1.1, 1.1])
    else:
        doc.add_paragraph("No phase data recorded.")

    doc.add_paragraph()

    # ── 5. LogGuard Probe Latency Timeline ─────────────────────────────────
    _add_heading(doc, "5. LogGuard Health Probe Latency (ms)")

    probes = data.get("logguard_all_probe_ms", [])
    if probes:
        # Summary stats
        avg_ms = sum(probes) / len(probes)
        max_ms = max(probes)
        min_ms = min(probes)
        p95_ms = sorted(probes)[int(len(probes) * 0.95)] if len(probes) >= 20 else max_ms

        _add_kv(doc, "Probes recorded", len(probes))
        _add_kv(doc, "Min latency",     f"{min_ms:.0f} ms")
        _add_kv(doc, "Avg latency",     f"{avg_ms:.0f} ms")
        _add_kv(doc, "95th percentile", f"{p95_ms:.0f} ms")
        _add_kv(doc, "Max latency",     f"{max_ms:.0f} ms",
                 value_color=COLOR_CRASH if max_ms > avg_ms * 5 else None)

        doc.add_paragraph()

        # Show probes in a compact table (max 40 rows to keep doc readable)
        display = probes if len(probes) <= 40 else probes[:20] + ["..."] + probes[-20:]
        headers = ["Probe #", "Response time (ms)"]
        rows    = []
        idx     = 1
        for v in display:
            if v == "...":
                rows.append(["...", "..."])
            else:
                rows.append([idx, f"{v:.0f}"])
                idx += 1
        _add_table(doc, headers, rows, col_widths=[1.0, 2.0])
    else:
        doc.add_paragraph("No probe data recorded.")

    doc.add_paragraph()

    # ── 6. Interpretation & Recommendations ───────────────────────────────
    _add_heading(doc, "6. Interpretation & Recommendations")

    total_logs = data.get("total_logs_accepted", 0)
    peak_rate  = data.get("peak_rate_total_logs_per_s", 0)
    baseline   = data.get("logguard_baseline_ms")
    final_ms   = data.get("logguard_avg_ms_at_end")

    bullets = []

    if crash:
        bullets.append(
            f"The server reached its limit at approximately {total_logs:,} total logs "
            f"(peak rate: {peak_rate} logs/s). Beyond this point the service became "
            "unavailable. This is the maximum sustained throughput for the current deployment."
        )
        bullets.append(
            "Recommended actions: scale up the anomaly-detection or Elasticsearch pods "
            "(increase K8s resource limits), add an Elasticsearch replica, or implement "
            "a queue/buffer (e.g. Kafka) between Fluent Bit and the backend."
        )
    else:
        bullets.append(
            "The server handled all test phases without crashing. "
            f"Maximum tested rate was {peak_rate} logs/s ({total_logs:,} total logs). "
            "The actual breaking point may be higher — re-run with increased MAX phase rate."
        )

    if baseline and final_ms:
        ratio = final_ms / baseline
        if ratio > 3:
            bullets.append(
                "LogGuard response time degraded significantly under load: "
                f"from {baseline:.0f}ms (baseline) to {final_ms:.0f}ms ({ratio:.1f}× slower). "
                "This indicates the backend is under memory or CPU pressure at peak rate."
            )
        else:
            bullets.append(
                f"Response time remained stable: {baseline:.0f}ms baseline → {final_ms:.0f}ms at peak "
                f"({ratio:.1f}× change). The server handled this load without significant latency increase."
            )

    bullets.append(
        "Post-test cleanup: delete the stress test log entries from Elasticsearch to avoid "
        "polluting real org data and inflating the warmup log_count. Use the Kibana admin "
        "panel or an ES delete-by-query filtered by the test timestamp range."
    )

    for b in bullets:
        para = doc.add_paragraph(style="List Bullet")
        run  = para.add_run(b)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Footer ──────────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run(
        "LogGuard Stress Test Report  •  "
        f"Generated {datetime.datetime.now(datetime.timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

    doc.save(out_path)
    print(f"  Word report saved: {out_path}")

# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Find JSON file
    if len(sys.argv) >= 2:
        json_path = sys.argv[1]
    else:
        # Auto-pick the most recent results file
        pattern = os.path.join(script_dir, "stress_test_results_*.json")
        files   = sorted(glob.glob(pattern))
        if not files:
            print("ERROR: No stress_test_results_*.json found.")
            print("Run stress_test.py first, then re-run this script.")
            sys.exit(1)
        json_path = files[-1]
        print(f"  Auto-selected: {json_path}")

    if not os.path.exists(json_path):
        print(f"ERROR: File not found: {json_path}")
        sys.exit(1)

    with open(json_path) as f:
        data = json.load(f)

    ts       = datetime.datetime.now(datetime.timezone.utc).strftime("%Y%m%d_%H%M%S")
    out_path = os.path.join(script_dir, f"stress_test_report_{ts}.docx")

    print(f"\n  Building Word report from: {os.path.basename(json_path)}")
    build_report(data, out_path)
    print("  Done.\n")


if __name__ == "__main__":
    main()
